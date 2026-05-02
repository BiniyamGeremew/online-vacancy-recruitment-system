"""
CV Text Extraction Service
Extracts text content from uploaded CV files (PDF and DOCX).
"""

import io
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document


def extract_text_from_cv(file):
    """
    Extract text from CV file.

    Args:
        file: Django FileField or file-like object

    Returns:
        str: Extracted plain text from the CV

    Raises:
        ValueError: If file format is unsupported or extraction fails
    """
    if not file:
        raise ValueError("No file provided")

    # Get file extension
    filename = getattr(file, 'name', str(file)).lower()

    try:
        if filename.endswith('.pdf'):
            return _extract_from_pdf(file)
        elif filename.endswith('.docx'):
            return _extract_from_docx(file)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")
    except Exception as e:
        raise ValueError(f"Failed to extract text from CV: {str(e)}")


def _extract_from_pdf(file):
    """Extract text from PDF file."""
    if hasattr(file, 'file'):
        file = file.file

    if hasattr(file, 'seek'):
        file.seek(0)

    if hasattr(file, 'read'):
        data = file.read()
        if isinstance(data, str):
            data = data.encode('utf-8')
        pdf_stream = io.BytesIO(data)
    else:
        raise ValueError("Unsupported PDF file object")

    text = extract_pdf_text(pdf_stream)

    if not text.strip():
        raise ValueError("No text content found in PDF")

    return text.strip()


def _extract_from_docx(file):
    """Extract text from DOCX file."""
    if hasattr(file, 'file'):
        file = file.file

    if hasattr(file, 'seek'):
        file.seek(0)

    if hasattr(file, 'read'):
        data = file.read()
        if isinstance(data, str):
            data = data.encode('utf-8')
        doc_stream = io.BytesIO(data)
    else:
        raise ValueError("Unsupported DOCX file object")

    doc = Document(doc_stream)

    # Extract text from all paragraphs
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    text = '\n'.join(text_parts)

    if not text.strip():
        raise ValueError("No text content found in DOCX")

    return text.strip()